import copy
import io
import json
import os
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseUpload

SCOPES = [
    "https://www.googleapis.com/auth/documents",
    "https://www.googleapis.com/auth/drive",
    "https://www.googleapis.com/auth/spreadsheets",
]

PROJECT_ROOT = Path(__file__).parent.parent
TOKEN_PATH = PROJECT_ROOT / "token.json"
SECRETS_PATH = PROJECT_ROOT / "client_secrets.json"
LOCAL_CONFIG_PATH = PROJECT_ROOT / "local_config.json"

SHEET_HEADERS = ["Date", "Company", "Title", "Seniority", "Match Score", "Job Post Link", "Tailored CV Link", "Submitted?"]


def _creds() -> Credentials:
    token_json_env = os.getenv("GOOGLE_TOKEN_JSON", "").strip()
    if token_json_env:
        creds = Credentials.from_authorized_user_info(json.loads(token_json_env), SCOPES)
        if not creds.valid and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        return creds

    if not SECRETS_PATH.exists():
        raise FileNotFoundError(
            f"client_secrets.json not found at {SECRETS_PATH}. "
            "Create OAuth 2.0 credentials (Desktop app) in Google Cloud Console and download them here."
        )
    creds = None
    if TOKEN_PATH.exists():
        creds = Credentials.from_authorized_user_file(str(TOKEN_PATH), SCOPES)
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            flow = InstalledAppFlow.from_client_secrets_file(str(SECRETS_PATH), SCOPES)
            creds = flow.run_local_server(port=0)
        TOKEN_PATH.write_text(creds.to_json())
    return creds


def _load_local_config() -> dict:
    if LOCAL_CONFIG_PATH.exists():
        return json.loads(LOCAL_CONFIG_PATH.read_text())
    return {}


def _save_local_config(data: dict) -> None:
    config = _load_local_config()
    config.update(data)
    LOCAL_CONFIG_PATH.write_text(json.dumps(config, indent=2))


def _get_or_create_company_folder(drive, company_name: str, parent_id: str = None) -> str:
    """Return the ID of a Drive folder named company_name, creating it if needed."""
    safe = company_name.replace("\\", "\\\\").replace("'", "\\'")
    parent_clause = f" and '{parent_id}' in parents" if parent_id else ""
    query = (
        f"name='{safe}' and mimeType='application/vnd.google-apps.folder'"
        f"{parent_clause} and trashed=false"
    )
    hits = drive.files().list(q=query, fields="files(id)", pageSize=1).execute().get("files", [])
    if hits:
        return hits[0]["id"]
    body = {"name": company_name, "mimeType": "application/vnd.google-apps.folder"}
    if parent_id:
        body["parents"] = [parent_id]
    return drive.files().create(body=body, fields="id").execute()["id"]


def _find_base_cv_id(drive) -> str:
    """Return the file ID of 'Tahel Tabacznik - Base CV'. Checks env var first, then searches by name."""
    base_cv_id = os.getenv("GOOGLE_CV_TEMPLATE_ID", "").strip()
    if base_cv_id:
        return base_cv_id

    output_folder_id = os.getenv("GOOGLE_OUTPUT_FOLDER_ID", "").strip()
    parent_clause = f" and '{output_folder_id}' in parents" if output_folder_id else ""
    query = (
        f"name='Tahel Tabacznik - Base CV' and mimeType='application/vnd.google-apps.document'"
        f"{parent_clause} and trashed=false"
    )
    hits = drive.files().list(q=query, fields="files(id)", pageSize=1).execute().get("files", [])
    if hits:
        return hits[0]["id"]

    # Widen search to all of Drive if not found in the output folder
    if output_folder_id:
        hits = drive.files().list(
            q="name='Tahel Tabacznik - Base CV' and mimeType='application/vnd.google-apps.document' and trashed=false",
            fields="files(id)", pageSize=1,
        ).execute().get("files", [])
        if hits:
            return hits[0]["id"]

    return ""


def read_base_cv_text() -> str:
    """Export the Base CV as plain text. Returns empty string if not found."""
    try:
        creds = _creds()
        drive = build("drive", "v3", credentials=creds)
        base_cv_id = _find_base_cv_id(drive)
        if not base_cv_id:
            return ""
        raw = drive.files().export(fileId=base_cv_id, mimeType="text/plain").execute()
        return raw.decode("utf-8") if isinstance(raw, bytes) else raw
    except Exception:
        return ""


def _set_para_text(para, new_text: str) -> None:
    """Replace all text in a paragraph preserving the first run's formatting (font, size, bold, etc.)."""
    runs = para.runs
    if not runs:
        para.add_run(new_text)
        return
    rpr = runs[0]._r.find(qn("w:rPr"))
    for r in para._p.findall(qn("w:r")):
        para._p.remove(r)
    new_r = OxmlElement("w:r")
    if rpr is not None:
        new_r.append(copy.deepcopy(rpr))
    new_t = OxmlElement("w:t")
    new_t.text = new_text
    new_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    new_r.append(new_t)
    para._p.append(new_r)


def _replace_bullets_in_doc(doc, company_marker: str, new_bullets: list) -> None:
    """Find bullet paragraphs after company_marker header and replace with new_bullets.
    Inserts or removes paragraphs to match the desired count."""
    paras = doc.paragraphs
    company_idx = next(
        (i for i, p in enumerate(paras) if p.text.strip().startswith(company_marker)), None
    )
    if company_idx is None:
        return
    job_title_idx = next(
        (i for i in range(company_idx + 1, len(paras)) if paras[i].text.strip()), None
    )
    if job_title_idx is None:
        return
    bullet_paras = []
    for i in range(job_title_idx + 1, len(paras)):
        if not paras[i].text.strip():
            break
        bullet_paras.append(paras[i])
    if not bullet_paras:
        return
    n_cur, n_new = len(bullet_paras), len(new_bullets)
    for i in range(min(n_cur, n_new)):
        _set_para_text(bullet_paras[i], new_bullets[i])
    if n_new > n_cur:
        ref_p = bullet_paras[0]._p
        insert_after = bullet_paras[-1]._p
        for i in range(n_cur, n_new):
            new_p = copy.deepcopy(ref_p)
            for r in new_p.findall(qn("w:r")):
                new_p.remove(r)
            rpr_src = ref_p.find(f".//{qn('w:r')}/{qn('w:rPr')}")
            new_r = OxmlElement("w:r")
            if rpr_src is not None:
                new_r.append(copy.deepcopy(rpr_src))
            new_t = OxmlElement("w:t")
            new_t.text = new_bullets[i]
            new_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            new_r.append(new_t)
            new_p.append(new_r)
            insert_after.addnext(new_p)
            insert_after = new_p
    elif n_new < n_cur:
        for para in bullet_paras[n_new:]:
            para._p.getparent().remove(para._p)


def _apply_cv_data(docx_bytes: bytes, cv_data: dict) -> bytes:
    """Apply structured cv_data to a .docx template, preserving all formatting. Returns modified bytes."""
    doc = Document(io.BytesIO(docx_bytes))
    paras = doc.paragraphs
    summary = cv_data.get("summary", "")
    if summary:
        for i, para in enumerate(paras):
            if para.text.strip() == "Professional Summary":
                for j in range(i + 1, len(paras)):
                    if paras[j].text.strip():
                        _set_para_text(paras[j], summary)
                        break
                break
    if cv_data.get("admaven_bullets"):
        _replace_bullets_in_doc(doc, "ADMAVEN", cv_data["admaven_bullets"])
    if cv_data.get("aa_financial_bullets"):
        _replace_bullets_in_doc(doc, "A. A. Financial", cv_data["aa_financial_bullets"])
    if cv_data.get("adore_bullets"):
        _replace_bullets_in_doc(doc, "Adore", cv_data["adore_bullets"])
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def create_tailored_cv_doc(cv_content, job_info: dict) -> str:
    """Create a tailored CV Google Doc preserving the base .docx formatting.
    cv_content: dict with summary + bullets (from tailor_cv). Returns Google Doc edit URL."""
    creds = _creds()
    docs = build("docs", "v1", credentials=creds)
    drive = build("drive", "v3", credentials=creds)

    company = job_info.get("company_name", "Company")
    file_name = "Tahel Tabacznik - CV"

    base_folder_id = os.getenv("GOOGLE_OUTPUT_FOLDER_ID", "").strip() or None
    company_folder_id = _get_or_create_company_folder(drive, company, base_folder_id)
    base_cv_id = _find_base_cv_id(drive)
    doc_id = None

    if base_cv_id and isinstance(cv_content, dict):
        try:
            docx_bytes = drive.files().export(
                fileId=base_cv_id,
                mimeType="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ).execute()
            modified_docx = _apply_cv_data(docx_bytes, cv_content)
            doc = drive.files().create(
                body={
                    "name": file_name,
                    "mimeType": "application/vnd.google-apps.document",
                    "parents": [company_folder_id],
                },
                media_body=MediaIoBaseUpload(
                    io.BytesIO(modified_docx),
                    mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                ),
                fields="id",
            ).execute()
            doc_id = doc["id"]
        except Exception:
            doc_id = None

    if doc_id is None:
        # Fallback: copy base CV Google Doc and replace text
        if base_cv_id:
            doc = drive.files().copy(
                fileId=base_cv_id,
                body={"name": file_name, "parents": [company_folder_id]},
            ).execute()
            doc_id = doc["id"]
        else:
            doc = docs.documents().create(body={"title": file_name}).execute()
            doc_id = doc["documentId"]
            drive.files().update(
                fileId=doc_id,
                addParents=company_folder_id,
                removeParents="root",
                fields="id, parents",
            ).execute()
        cv_text = "\n".join(
            [cv_content.get("summary", "")] +
            cv_content.get("admaven_bullets", []) +
            cv_content.get("aa_financial_bullets", []) +
            cv_content.get("adore_bullets", [])
        ) if isinstance(cv_content, dict) else str(cv_content)
        if cv_text.strip():
            existing = docs.documents().get(documentId=doc_id).execute()
            end_index = existing["body"]["content"][-1]["endIndex"] - 1
            requests = []
            if end_index > 1:
                requests.append({"deleteContentRange": {"range": {"startIndex": 1, "endIndex": end_index}}})
            requests.append({"insertText": {"location": {"index": 1}, "text": cv_text}})
            docs.documents().batchUpdate(documentId=doc_id, body={"requests": requests}).execute()

    # Export as PDF and upload to the same company folder
    pdf_bytes = drive.files().export(fileId=doc_id, mimeType="application/pdf").execute()
    drive.files().create(
        body={"name": file_name, "parents": [company_folder_id]},
        media_body=MediaIoBaseUpload(io.BytesIO(pdf_bytes), mimetype="application/pdf"),
        fields="id",
    ).execute()

    # Share the doc so the edit link in the UI works
    drive.permissions().create(
        fileId=doc_id,
        body={"type": "anyone", "role": "reader"},
    ).execute()

    cv_url = f"https://docs.google.com/document/d/{doc_id}/edit"
    folder_url = f"https://drive.google.com/drive/folders/{company_folder_id}"
    return cv_url, folder_url


def log_to_sheet(job_info: dict, match_score: int, cv_url: str, job_url: str) -> str:
    """Append a row to the tracking sheet. Creates the sheet on first use. Returns sheet URL."""
    creds = _creds()
    sheets = build("sheets", "v4", credentials=creds)
    drive = build("drive", "v3", credentials=creds)

    sheet_id = os.getenv("GOOGLE_SHEET_ID", "").strip() or _load_local_config().get("sheet_id", "")

    if not sheet_id:
        sheet_id = _create_tracking_sheet(sheets, drive)
        _save_local_config({"sheet_id": sheet_id})
        print(f"\n[CV Updater] Tracking sheet created. Add to .env:\nGOOGLE_SHEET_ID={sheet_id}\n")

    row = [
        datetime.now().strftime("%Y-%m-%d"),
        job_info.get("company_name", ""),
        job_info.get("job_title", ""),
        job_info.get("seniority", ""),
        match_score,
        job_url,
        cv_url,
        False,  # Submitted?
    ]
    result = sheets.spreadsheets().values().append(
        spreadsheetId=sheet_id,
        range="Sheet1!A:H",
        valueInputOption="RAW",
        body={"values": [row]},
    ).execute()

    # Apply checkbox data validation to the Submitted? cell (column H)
    updated_range = result.get("updates", {}).get("updatedRange", "")
    row_match = re.search(r":H(\d+)$", updated_range)
    if row_match:
        row_idx = int(row_match.group(1)) - 1  # convert to 0-based
        meta = sheets.spreadsheets().get(spreadsheetId=sheet_id, fields="sheets.properties").execute()
        grid_id = meta["sheets"][0]["properties"]["sheetId"]
        sheets.spreadsheets().batchUpdate(
            spreadsheetId=sheet_id,
            body={"requests": [{
                "setDataValidation": {
                    "range": {
                        "sheetId": grid_id,
                        "startRowIndex": row_idx,
                        "endRowIndex": row_idx + 1,
                        "startColumnIndex": 7,
                        "endColumnIndex": 8,
                    },
                    "rule": {"condition": {"type": "BOOLEAN"}, "showCustomUi": True},
                }
            }]}
        ).execute()

    return f"https://docs.google.com/spreadsheets/d/{sheet_id}"


def check_duplicate(company_name: str, job_title: str):
    """Check if this company+title combo already exists in the tracker. Returns existing row data or None."""
    try:
        creds = _creds()
        sheets = build("sheets", "v4", credentials=creds)
        sheet_id = os.getenv("GOOGLE_SHEET_ID", "").strip() or _load_local_config().get("sheet_id", "")
        if not sheet_id:
            return None

        result = sheets.spreadsheets().values().get(
            spreadsheetId=sheet_id,
            range="Sheet1!A:H",
        ).execute()
        rows = result.get("values", [])
        if len(rows) <= 1:
            return None

        company_lower = company_name.lower().strip()
        title_lower = job_title.lower().strip()

        for row in rows[1:]:
            if len(row) < 3:
                continue
            if row[1].lower().strip() == company_lower and row[2].lower().strip() == title_lower:
                return {
                    "date": row[0] if len(row) > 0 else "",
                    "company": row[1] if len(row) > 1 else "",
                    "title": row[2] if len(row) > 2 else "",
                    "cv_url": row[6] if len(row) > 6 else "",
                }
        return None
    except Exception:
        return None  # Never block a run due to a failed check


def _create_tracking_sheet(sheets, drive) -> str:
    header_cells = [{"userEnteredValue": {"stringValue": h}} for h in SHEET_HEADERS]
    spreadsheet = sheets.spreadsheets().create(body={
        "properties": {"title": "Tahel — Job Applications Tracker"},
        "sheets": [{
            "properties": {"title": "Sheet1"},
            "data": [{"rowData": [{"values": header_cells}]}],
        }],
    }).execute()
    sheet_id = spreadsheet["spreadsheetId"]
    drive.permissions().create(
        fileId=sheet_id,
        body={"type": "anyone", "role": "writer"},
    ).execute()
    return sheet_id
